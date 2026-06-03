import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from google import genai
from google.genai import types
from PIL import Image
import pandas as pd
import io
import json
import re

# -------------------------------
# Page Setup
# -------------------------------
st.set_page_config(
    page_title="Document / PDF / Visual Scraping AI Agent",
    layout="wide"
)

st.title("📄 Document / PDF / Visual Scraping AI Agent")
st.write(
    "Upload PDF, DOCX, JPG, or PNG files. The app can extract document details, "
    "commercial rates, visual/scanned information, and perform complex calculations."
)

# -------------------------------
# Gemini API Key
# -------------------------------
try:
    api_key = st.secrets["GEMINI_API_KEY"]
except Exception:
    st.error("Gemini API key is missing. Please add GEMINI_API_KEY in Streamlit Secrets.")
    st.stop()

client = genai.Client(api_key=api_key)

# -------------------------------
# File Upload
# -------------------------------
uploaded_file = st.file_uploader(
    "Upload your file",
    type=["pdf", "docx", "jpg", "jpeg", "png"]
)

extraction_type = st.selectbox(
    "Select extraction type",
    [
        "Ask Any Question / Complex Calculation",
        "Commercial Rate Extraction",
        "Contract Details",
        "Billing Details",
        "Invoice Details",
        "Visual / Scanned Document Details",
        "Risk / Missing Information",
        "Custom Question"
    ]
)

user_question = ""

if extraction_type in ["Ask Any Question / Complex Calculation", "Custom Question"]:
    user_question = st.text_area(
        "Ask your question",
        placeholder=(
            "Example: Calculate annual plot rent for 50,000 sqm at AED 80 per sqm per annum. "
            "Or calculate 5-year revenue with 2.5% escalation."
        )
    )

analyze_visual = st.checkbox(
    "Analyze visual/scanned content also",
    value=True
)

max_pages = st.slider(
    "Maximum PDF pages to visually analyze",
    min_value=1,
    max_value=15,
    value=5
)

# -------------------------------
# PDF Text Extraction
# -------------------------------
def extract_pdf_text(file_bytes):
    text = ""

    try:
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")

        for page_num, page in enumerate(pdf_document, start=1):
            page_text = page.get_text()
            text += f"\n\n--- Page {page_num} ---\n"
            text += page_text

        return text

    except Exception as e:
        return f"PDF text extraction failed: {str(e)}"


# -------------------------------
# Convert PDF Pages to Images
# -------------------------------
def convert_pdf_pages_to_images(file_bytes, max_pages=5):
    images = []

    try:
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = min(len(pdf_document), max_pages)

        for page_index in range(total_pages):
            page = pdf_document[page_index]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            img_bytes = pix.tobytes("png")

            image = Image.open(io.BytesIO(img_bytes)).convert("RGB")
            images.append(image)

        return images

    except Exception as e:
        st.warning(f"PDF visual conversion failed: {str(e)}")
        return []


# -------------------------------
# DOCX Text Extraction
# -------------------------------
def extract_docx_text(file):
    text = ""

    try:
        doc = Document(file)

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text += " | ".join(cells) + "\n"

        return text

    except Exception as e:
        return f"DOCX extraction failed: {str(e)}"


# -------------------------------
# Question Builder
# -------------------------------
def get_question(extraction_type, user_question):

    if extraction_type == "Ask Any Question / Complex Calculation":
        return f"""
Answer the user's question based on the uploaded document.

User question:
{user_question}

You must understand the document and perform calculations if required.

Calculation rules:
1. Extract values from the document first.
2. Use exact text from the document as evidence.
3. If the user gives additional values in the question, use them also.
4. Do not guess missing values.
5. If any value is missing, mention it clearly.
6. Show formula used.
7. Show step-by-step working.
8. Give final answer clearly.
9. Handle complex commercial calculations such as:
   - Area × rate
   - Linear meter × daily rate × number of days
   - Annual, monthly, quarterly, daily pro-rata calculations
   - Escalation year-on-year
   - Multi-year revenue
   - Tax, discount, penalty, minimum commitment
   - Contract billing period calculation
   - Total payable amount
10. If the document says “AED 80 per m² per annum”, understand it as AED 80 per square meter per year.
11. If the document says “AED 26 per linear meter per day”, understand it as AED 26 per linear meter per day.
"""

    elif extraction_type == "Commercial Rate Extraction":
        return """
Extract all commercial rates, charges, tariffs, rent, fees, service charges, quay wall charges, utility charges,
waste management charges, storage charges, plot charges, land lease rates, escalation rates, and billing rates
from the document.

For every rate found, provide:
1. Charge type / rate name
2. Exact text from the document without changing the wording
3. Currency
4. Rate amount
5. Unit of measurement, such as per m², per sqm, per square meter, per linear meter, per day, per annum
6. Billing frequency
7. Whether the rate is fixed, variable, tariff-based, prevailing-rate based, or escalation-based
8. Clear business meaning in simple words
9. Page number or section reference
10. Missing or unclear details
"""

    elif extraction_type == "Contract Details":
        return """
Extract the following contract details:
1. Customer / Project name
2. Effective date
3. Contract term
4. Handover date
5. Commercial operation date
6. Total area
7. Commercial rates
8. Escalation clause
9. Payment terms
10. Billing start date
11. Important obligations
12. Missing or unclear commercial terms

For commercial rates, include exact text, amount, currency, unit, frequency, and business meaning.
"""

    elif extraction_type == "Billing Details":
        return """
Extract the following billing details:
1. Billing start date
2. Billing frequency
3. Billing rate
4. Currency
5. Quantity / area / volume
6. Payment terms
7. Tax or additional charges
8. Escalation rule
9. Billing conditions
10. Any unclear billing terms
"""

    elif extraction_type == "Invoice Details":
        return """
Extract the following invoice details:
1. Invoice number
2. Invoice date
3. Customer name
4. Amount
5. Currency
6. PO number
7. Job / shipment reference
8. Tax amount
9. Total invoice value
10. Any discrepancy or missing information
"""

    elif extraction_type == "Visual / Scanned Document Details":
        return """
Read the visual or scanned document carefully and extract:
1. All visible important text
2. Tables and values
3. Dates
4. Amounts
5. Commercial rates
6. Customer / company names
7. Contract or invoice references
8. Signatures, stamps, handwritten notes if visible
9. Any unclear or unreadable parts
"""

    elif extraction_type == "Risk / Missing Information":
        return """
Review the document and identify:
1. Missing information
2. Unclear dates
3. Unclear commercial terms
4. Missing rates
5. Missing payment terms
6. Missing billing start date
7. Assumptions
8. Risk points
9. Clauses requiring manual review
"""

    else:
        return user_question


# -------------------------------
# Prompt Builder
# -------------------------------
def build_prompt(document_text, question):
    return f"""
You are an expert document scraping AI agent for contracts, invoices, billing documents, and commercial documents.

Your job is to:
1. Read and understand the document.
2. Extract exact commercial and contract details.
3. Perform calculations when asked.
4. Show exact document text as evidence.
5. Avoid guessing missing values.

The document may contain:
- Normal text
- Tables
- Scanned pages
- Images
- Stamps
- Signatures
- Commercial clauses
- Billing rates
- Tariff details
- Contract terms

User request:
{question}

Extracted text from document:
{document_text}

Return only valid JSON in the following format:

{{
  "summary": "Short summary of the document and the answer",
  "extracted_details": [
    {{
      "field": "Charge type or field name",
      "exact_text_from_document": "Copy the exact wording from the document",
      "currency": "Currency, for example AED",
      "rate_amount": "Amount only, for example 80.00. If no amount, write Not fixed",
      "unit": "Unit, for example per m², per sqm, per linear meter, per day",
      "frequency": "Frequency, for example per annum, per day, monthly, one-time, as applicable",
      "rate_type": "Fixed / Variable / Tariff-based / Prevailing rate / Escalation-based / Not applicable",
      "business_meaning": "Explain the meaning in simple business language",
      "reference": "Page number, section, clause, table, or visual reference if available"
    }}
  ],
  "calculation_results": [
    {{
      "calculation_name": "Name of the calculation",
      "inputs_used": "List the values used from the document or user question",
      "formula": "Formula used",
      "step_by_step_working": "Show the calculation steps clearly",
      "final_result": "Final calculated result",
      "currency": "Currency if applicable",
      "unit": "Unit if applicable",
      "notes": "Mention assumptions or missing inputs"
    }}
  ],
  "missing_or_unclear_details": [
    "Mention anything missing, unclear, unreadable, or assumed"
  ]
}}

Important:
- Return JSON only.
- Do not add markdown.
- Do not add explanation outside JSON.
- Do not guess values that are not present.
- Preserve exact commercial text from the document.
- For calculations, always show formula and working.
"""


# -------------------------------
# Gemini Text Analysis
# -------------------------------
def ask_gemini_text(document_text, question):
    prompt = build_prompt(document_text, question)

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json"
        )
    )

    return response.text


# -------------------------------
# Gemini Visual Analysis
# -------------------------------
def ask_gemini_visual(document_text, question, images):
    prompt = build_prompt(document_text, question)

    contents = [prompt]

    for image in images:
        contents.append(image)

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=contents,
        config=types.GenerateContentConfig(
            response_mime_type="application/json"
        )
    )

    return response.text


# -------------------------------
# JSON Cleaner
# -------------------------------
def clean_json_response(ai_result):
    cleaned = ai_result.strip()
    cleaned = cleaned.replace("```json", "")
    cleaned = cleaned.replace("```", "")
    cleaned = cleaned.strip()

    try:
        return json.loads(cleaned)
    except Exception:
        pass

    match = re.search(r"\{.*\}", cleaned, re.DOTALL)

    if match:
        return json.loads(match.group(0))

    raise ValueError("Could not parse AI response as JSON.")


# -------------------------------
# Excel Creator
# -------------------------------
def create_excel_file(details_df, calculation_df):
    excel_buffer = io.BytesIO()

    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        if not details_df.empty:
            details_df.to_excel(writer, index=False, sheet_name="Extracted Details")

        if not calculation_df.empty:
            calculation_df.to_excel(writer, index=False, sheet_name="Calculations")

    excel_buffer.seek(0)
    return excel_buffer


# -------------------------------
# Show Result
# -------------------------------
def show_result(ai_result):
    try:
        result_json = clean_json_response(ai_result)

        st.subheader("Summary")
        st.write(result_json.get("summary", ""))

        details = result_json.get("extracted_details", [])
        calculations = result_json.get("calculation_results", [])

        details_df = pd.DataFrame(details)
        calculation_df = pd.DataFrame(calculations)

        if not details_df.empty:
            st.subheader("Extracted Details Table")
            st.dataframe(details_df, use_container_width=True)

        if not calculation_df.empty:
            st.subheader("Calculation Results")
            st.dataframe(calculation_df, use_container_width=True)

        if not details_df.empty or not calculation_df.empty:
            excel_file = create_excel_file(details_df, calculation_df)

            st.download_button(
                label="Download Excel Report",
                data=excel_file.getvalue(),
                file_name="document_ai_extraction_and_calculation.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        missing = result_json.get("missing_or_unclear_details", [])

        if missing:
            st.subheader("Missing / Unclear Details")
            for item in missing:
                st.warning(item)

        with st.expander("View Raw JSON"):
            st.json(result_json)

    except Exception as e:
        st.subheader("AI Extracted Result")
        st.write(ai_result)
        st.info("Excel conversion was skipped because the AI response was not valid JSON.")
        st.write(str(e))


# -------------------------------
# Main Button
# -------------------------------
if st.button("Extract / Calculate"):

    if not uploaded_file:
        st.error("Please upload a file.")

    elif extraction_type in ["Ask Any Question / Complex Calculation", "Custom Question"] and not user_question:
        st.error("Please enter your question.")

    else:
        try:
            question = get_question(extraction_type, user_question)
            file_name = uploaded_file.name.lower()
            file_bytes = uploaded_file.getvalue()

            document_text = ""
            images = []

            with st.spinner("Reading file, understanding document, and preparing answer..."):

                if file_name.endswith(".pdf"):
                    document_text = extract_pdf_text(file_bytes)

                    if analyze_visual:
                        images = convert_pdf_pages_to_images(file_bytes, max_pages=max_pages)

                    if len(document_text.strip()) < 50:
                        st.warning(
                            "Very little text was extracted. This may be a scanned PDF. "
                            "Visual analysis will be used if enabled."
                        )

                    if analyze_visual and images:
                        ai_result = ask_gemini_visual(document_text, question, images)
                    else:
                        ai_result = ask_gemini_text(document_text, question)

                elif file_name.endswith(".docx"):
                    document_text = extract_docx_text(uploaded_file)
                    ai_result = ask_gemini_text(document_text, question)

                elif (
                    file_name.endswith(".jpg")
                    or file_name.endswith(".jpeg")
                    or file_name.endswith(".png")
                ):
                    image = Image.open(io.BytesIO(file_bytes)).convert("RGB")

                    st.image(
                        image,
                        caption="Uploaded Image",
                        use_container_width=True
                    )

                    images = [image]
                    document_text = (
                        "The uploaded file is an image. Read all visible text, "
                        "tables, rates, stamps, signatures, and visual information."
                    )

                    ai_result = ask_gemini_visual(document_text, question, images)

                else:
                    st.error("Unsupported file type.")
                    st.stop()

            show_result(ai_result)

        except Exception as e:
            st.error("Something went wrong.")
            st.write(str(e))
